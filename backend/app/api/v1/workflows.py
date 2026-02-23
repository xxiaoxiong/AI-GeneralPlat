import io
import re
import time
import json
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from typing import Optional, Any
from pydantic import BaseModel

from app.core.database import get_db
from app.models.user import User
from app.models.workflow import Workflow, WorkflowExecution
from app.schemas.workflow import (
    WorkflowCreate, WorkflowUpdate, WorkflowOut,
    WorkflowExecuteRequest, WorkflowExecutionOut,
)
from app.schemas.common import PageResponse
from app.api.deps import get_current_user
from app.services.workflow_service import WorkflowEngine
from app.services.audit import log_action

router = APIRouter(prefix="/workflows", tags=["流程编排"])


def _safe_json(obj: Any) -> Any:
    """Deep-serialize to plain Python types via JSON round-trip, stripping non-serializable objects."""
    try:
        return json.loads(json.dumps(obj, default=str))
    except Exception:
        return {}


@router.get("", response_model=PageResponse[WorkflowOut])
async def list_workflows(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    keyword: Optional[str] = None,
    category: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    query = select(Workflow).where(Workflow.owner_id == current_user.id)
    if keyword:
        query = query.where(Workflow.name.contains(keyword))
    if category:
        query = query.where(Workflow.category == category)

    total_result = await db.execute(select(func.count()).select_from(query.subquery()))
    total = total_result.scalar()

    query = query.offset((page - 1) * page_size).limit(page_size).order_by(Workflow.id.desc())
    result = await db.execute(query)
    workflows = result.scalars().all()

    return PageResponse(
        items=[WorkflowOut.model_validate(w) for w in workflows],
        total=total,
        page=page,
        page_size=page_size,
        total_pages=(total + page_size - 1) // page_size,
    )


@router.post("", response_model=WorkflowOut)
async def create_workflow(
    body: WorkflowCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    workflow = Workflow(**body.model_dump(), owner_id=current_user.id)
    db.add(workflow)
    await db.flush()
    await log_action(db, current_user.id, current_user.username, "create_workflow", "workflow",
                     resource_id=str(workflow.id))
    return WorkflowOut.model_validate(workflow)


@router.get("/{workflow_id}", response_model=WorkflowOut)
async def get_workflow(
    workflow_id: int,
    db: AsyncSession = Depends(get_db),
    _: User = Depends(get_current_user),
):
    result = await db.execute(select(Workflow).where(Workflow.id == workflow_id))
    workflow = result.scalar_one_or_none()
    if not workflow:
        raise HTTPException(status_code=404, detail="流程不存在")
    return WorkflowOut.model_validate(workflow)


@router.put("/{workflow_id}", response_model=WorkflowOut)
async def update_workflow(
    workflow_id: int,
    body: WorkflowUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Workflow).where(Workflow.id == workflow_id))
    workflow = result.scalar_one_or_none()
    if not workflow:
        raise HTTPException(status_code=404, detail="流程不存在")
    for field, value in body.model_dump(exclude_none=True).items():
        setattr(workflow, field, value)
    if body.definition is not None:
        workflow.version += 1
    db.add(workflow)
    return WorkflowOut.model_validate(workflow)


@router.delete("/{workflow_id}")
async def delete_workflow(
    workflow_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Workflow).where(Workflow.id == workflow_id))
    workflow = result.scalar_one_or_none()
    if not workflow:
        raise HTTPException(status_code=404, detail="流程不存在")
    await db.delete(workflow)
    return {"message": "删除成功"}


@router.post("/{workflow_id}/execute", response_model=WorkflowExecutionOut)
async def execute_workflow(
    workflow_id: int,
    body: WorkflowExecuteRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Workflow).where(Workflow.id == workflow_id))
    workflow = result.scalar_one_or_none()
    if not workflow:
        raise HTTPException(status_code=404, detail="流程不存在")
    if not workflow.is_active:
        raise HTTPException(status_code=400, detail="流程未启用")

    execution = WorkflowExecution(
        workflow_id=workflow_id,
        status="running",
        trigger_by=current_user.id,
        input_data=body.input_data,
    )
    db.add(execution)
    await db.flush()

    start_time = time.time()
    try:
        output, node_results = await WorkflowEngine.execute(workflow.definition, body.input_data or {}, db)
        execution.status = "success"
        execution.output_data = _safe_json(output)
        execution.node_results = _safe_json(node_results)
    except Exception as e:
        execution.status = "failed"
        execution.error_msg = str(e)

    execution.duration_ms = int((time.time() - start_time) * 1000)
    db.add(execution)
    await db.flush()
    await log_action(db, current_user.id, current_user.username, "execute_workflow", "workflow",
                     resource_id=str(workflow_id))
    await db.commit()
    # Build response from plain data to avoid circular reference via relationship
    return WorkflowExecutionOut(
        id=execution.id,
        workflow_id=execution.workflow_id,
        status=execution.status,
        input_data=execution.input_data,
        output_data=execution.output_data,
        node_results=execution.node_results,
        error_msg=execution.error_msg,
        duration_ms=execution.duration_ms,
        created_at=execution.created_at,
    )


@router.get("/{workflow_id}/executions", response_model=PageResponse[WorkflowExecutionOut])
async def list_executions(
    workflow_id: int,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: AsyncSession = Depends(get_db),
    _: User = Depends(get_current_user),
):
    query = select(WorkflowExecution).where(WorkflowExecution.workflow_id == workflow_id)
    total_result = await db.execute(select(func.count()).select_from(query.subquery()))
    total = total_result.scalar()

    query = query.offset((page - 1) * page_size).limit(page_size).order_by(WorkflowExecution.id.desc())
    result = await db.execute(query)
    executions = result.scalars().all()

    return PageResponse(
        items=[WorkflowExecutionOut.model_validate(e) for e in executions],
        total=total,
        page=page,
        page_size=page_size,
        total_pages=(total + page_size - 1) // page_size,
    )


class ExportDocRequest(BaseModel):
    content: str
    title: str = "工作流输出文档"
    filename: str = "output"


@router.post("/{workflow_id}/export-doc")
async def export_doc(
    workflow_id: int,
    body: ExportDocRequest,
    _: User = Depends(get_current_user),
):
    """Generate a Word document from workflow output content and return as download."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Title
    title_para = doc.add_heading(body.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown-like content into docx structure
    lines = body.content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('### '):
            p = doc.add_heading('', level=3)
            _add_inline_formatting(p, _strip_md_marks(line[4:].strip()))
        elif line.startswith('## '):
            p = doc.add_heading('', level=2)
            _add_inline_formatting(p, _strip_md_marks(line[3:].strip()))
        elif line.startswith('# '):
            p = doc.add_heading('', level=1)
            _add_inline_formatting(p, _strip_md_marks(line[2:].strip()))
        # Horizontal rule
        elif line.strip() in ('---', '***', '___'):
            doc.add_paragraph('─' * 40)
        # Unordered list
        elif re.match(r'^[-*+] ', line):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatting(p, line[2:].strip())
        # Ordered list
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line).strip()
            _add_inline_formatting(p, text)
        # Blockquote
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph('\n'.join(code_lines))
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        # Empty line → paragraph break
        elif line.strip() == '':
            pass
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, line)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from urllib.parse import quote
    safe_ascii = re.sub(r'[^\x00-\x7f]', '_', body.filename) or 'output'
    encoded_name = quote(body.filename or 'output', safe='')
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": (
                f'attachment; filename="{safe_ascii}.docx"; '
                f"filename*=UTF-8''{encoded_name}.docx"
            )
        },
    )


def _strip_md_marks(text: str) -> str:
    """Remove surrounding ** or * from a string (e.g. heading that is fully bolded)."""
    text = text.strip()
    if text.startswith('**') and text.endswith('**'):
        text = text[2:-2]
    elif text.startswith('*') and text.endswith('*'):
        text = text[1:-1]
    return text


def _add_inline_formatting(paragraph, text: str):
    """Parse **bold**, *italic*, `code` inline markdown and add runs."""
    from docx.shared import Pt, RGBColor
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        if m.group(1).startswith('**'):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(1).startswith('*'):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(1).startswith('`'):
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])
